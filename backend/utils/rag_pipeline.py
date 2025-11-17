"""
Helper utilities for building the retrieval-augmented generation (RAG) chain
that powers the Resume Assistant backend.

The functions in this module ingest resume materials from disk, chunk them into
retrieval-friendly segments, and wire up LangChain components to produce a
simple question-answering chain.

To keep startup latency low, the FAISS vector store and its metadata are
persisted to disk. Subsequent boots reuse the cached embeddings unless any of
the source resume files change.
"""

import json
import logging
import os
from pathlib import Path
from typing import Dict, List

from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate

logger = logging.getLogger(__name__)

VECTOR_CACHE_ENV_VAR = "RAG_FAISS_CACHE_DIR"
VECTOR_CACHE_DIR = Path(
    os.getenv(
        VECTOR_CACHE_ENV_VAR,
        Path(__file__).resolve().parents[1] / "vector_cache",
    )
).expanduser()
VECTOR_CACHE_METADATA_FILE = "metadata.json"
SUPPORTED_EXTENSIONS = (".txt", ".docx")

load_dotenv()

def _resolve_docs_dir():
    """Return the absolute path to the resumeMaterial directory."""

    return os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "resumeMaterial")
    )


def _load_text_file(path):
    """Load a UTF-8 text file into a LangChain Document with metadata."""

    with open(path, "r", encoding="utf-8") as infile:
        content = infile.read()
    return Document(page_content=content, metadata={"source": path})


def _load_docx_file(path):
    """Extract paragraphs from a .docx file and return a Document."""

    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)
    return Document(page_content=content, metadata={"source": path})


def _split_documents(documents, chunk_size=1000, chunk_overlap=150):
    """Split Documents into overlapping chunks suitable for retrieval."""

    split_docs = []
    for doc in documents:
        text = doc.page_content or ""
        metadata = dict(doc.metadata or {})
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk_text = text[start:end]
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_start"] = start
            chunk_metadata["chunk_end"] = end
            split_docs.append(Document(page_content=chunk_text, metadata=chunk_metadata))
            if end == len(text):
                break
            start = max(0, end - chunk_overlap)
    return split_docs


def _list_resume_files(docs_dir: Path) -> List[Path]:
    """Return all supported resume material files under docs_dir."""

    docs_dir = Path(docs_dir)
    files: List[Path] = []
    for path in docs_dir.rglob("*"):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return sorted(files)


def _build_snapshot(files: List[Path]) -> Dict[str, Dict[str, int]]:
    """Create a filesystem snapshot used to validate cached vector stores."""

    snapshot: Dict[str, Dict[str, int]] = {}
    for file_path in files:
        stats = file_path.stat()
        snapshot[str(file_path)] = {
            "mtime": int(stats.st_mtime),
            "size": stats.st_size,
        }
    return snapshot


def _load_cached_vector_store(cache_dir: Path, embeddings, snapshot: Dict[str, Dict[str, int]]):
    """Attempt to load a cached FAISS vector store that matches the snapshot."""

    metadata_path = cache_dir / VECTOR_CACHE_METADATA_FILE
    if not metadata_path.exists():
        return None

    try:
        stored_metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        logger.warning("Vector cache metadata missing or invalid; rebuilding index.")
        return None

    if stored_metadata.get("snapshot") != snapshot:
        logger.info("Resume materials changed since last cache; rebuilding index.")
        return None

    try:
        return FAISS.load_local(
            str(cache_dir),
            embeddings,
            allow_dangerous_deserialization=True,
        )
    except TypeError:
        # Backwards compatibility with older LangChain versions.
        return FAISS.load_local(str(cache_dir), embeddings)
    except Exception as exc:
        logger.warning("Failed to load cached vector store: %s", exc)
        return None


def _persist_vector_store(cache_dir: Path, vector_store, snapshot: Dict[str, Dict[str, int]]):
    """Save the FAISS index and associated snapshot metadata to disk."""

    cache_dir.mkdir(parents=True, exist_ok=True)
    vector_store.save_local(str(cache_dir))
    metadata_payload = {
        "snapshot": snapshot,
    }
    metadata_path = cache_dir / VECTOR_CACHE_METADATA_FILE
    metadata_path.write_text(json.dumps(metadata_payload, indent=2), encoding="utf-8")
    logger.info("Persisted FAISS index to %s", cache_dir)


def _load_documents(files: List[Path]) -> List[Document]:
    """Load supported resume files into LangChain Documents."""

    documents: List[Document] = []
    for file_path in files:
        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            documents.append(_load_text_file(file_path))
        elif suffix == ".docx":
            documents.append(_load_docx_file(file_path))
    return documents


def build_rag_chain():
    """
    Build and return a thin wrapper around a LangChain runnable that can
    answer questions about the ingested resume materials.
    """

    docs_dir = Path(_resolve_docs_dir())
    resume_files = _list_resume_files(docs_dir)

    if not resume_files:
        raise ValueError(f"No documents were loaded from {docs_dir}")

    snapshot = _build_snapshot(resume_files)

    embeddings = OpenAIEmbeddings()
    cache_dir = VECTOR_CACHE_DIR
    vector_store = _load_cached_vector_store(cache_dir, embeddings, snapshot)

    if vector_store is None:
        all_docs = _load_documents(resume_files)
        split_docs = _split_documents(all_docs, chunk_size=1000, chunk_overlap=150)
        vector_store = FAISS.from_documents(split_docs, embeddings)
        _persist_vector_store(cache_dir, vector_store, snapshot)
    else:
        logger.info("Loaded FAISS vector store from cache: %s", cache_dir)

    retriever = vector_store.as_retriever(search_kwargs={"k": 3})

    model_name = os.getenv("OPENAI_LLM_MODEL")
    temperature = 0
    llm = ChatOpenAI(model=model_name, temperature=temperature)
    prompt = ChatPromptTemplate.from_messages([
        (
            "system",
            (
                "You are an AI Resume Assistant representing Huang Jiashu, a Software Engineer specializing in "
                "backend development, Oracle APEX systems, and AI-driven applications. Huang Jiashu is passionate about "
                "machine learning and artificial intelligence.\n\n"
                "Your role is to help users understand his experience, projects, and technical skills.\n\n"
                "Primarily base your answers on the provided resume context. You may add light, positive phrasing "
                "to give a good professional impression of Jiashu, but do not invent specific facts that are not "
                "supported by the context.\n\n"
                "If the context does not contain the information needed to answer a question, respond with: "
                "\"I'm not sure, please refer to the resume for details.\"\n\n"
                "Keep responses concise, factual, and professional."
            ),
        ),
        (
            "human",
            (
                "Context:\n{context}\n\n"
                "User Question: {question}\n\n"
                "Respond naturally as if you're the AI version of Huang Jiashu introducing his experience."
            ),
        ),
    ])

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    extract_query = RunnableLambda(lambda inputs: inputs["query"])
    context_builder = extract_query | retriever | RunnableLambda(format_docs)
    runnable = (
        {"context": context_builder, "question": extract_query}
        | prompt
        | llm
        | StrOutputParser()
    )

    class SimpleQAChain:
        """Compatibility wrapper that returns dicts instead of raw strings."""

        def __init__(self, inner):
            self._inner = inner

        def invoke(self, inputs, **kwargs):
            result = self._inner.invoke(inputs, **kwargs)
            return {"result": result}

    return SimpleQAChain(runnable)
